"""
generate_test_report.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Taruh file ini di root project:
  public_speaking_trainer_with_GEMINI/generate_test_report.py

Cara pakai:
  python generate_test_report.py

Output:
  tests/laporan_pengujian.docx
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import subprocess
import sys
import os
import re
from datetime import datetime
from pathlib import Path


# ── Cek & install docx jika belum ada ─────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("📦 Menginstall python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════
# STEP 1 — JALANKAN PYTEST & TANGKAP OUTPUT
# ══════════════════════════════════════════════════════════════

def run_pytest():
    """Jalankan pytest dan kembalikan output sebagai string."""
    print("🧪 Menjalankan pytest...")
    result = subprocess.run(
        [sys.executable, "-m", "pytest", "tests/", "-v", "--tb=short", "--no-header"],
        capture_output=True,
        text=True,
        encoding="utf-8",
        cwd=Path(__file__).parent
    )
    output = result.stdout + result.stderr
    print("✅ Pytest selesai.")
    return output, result.returncode


# ══════════════════════════════════════════════════════════════
# STEP 2 — PARSE OUTPUT PYTEST
# ══════════════════════════════════════════════════════════════

def parse_pytest_output(output: str) -> dict:
    """
    Parse raw pytest output menjadi struktur data terorganisir.
    Kembalikan dict berisi list test results dan summary.
    """
    lines  = output.splitlines()
    tests  = []
    errors = []

    # Regex untuk baris hasil test
    # Contoh: tests/test_register.py::TestClass::test_func PASSED [ 12%]
    test_pattern = re.compile(
        r'^(tests[\\/][^\s]+\.py)::([^\s]+)\s+(PASSED|FAILED|ERROR|SKIPPED|XFAIL|XPASS)',
        re.IGNORECASE
    )
    # Regex untuk baris summary
    summary_pattern = re.compile(
        r'=+\s+(\d+\s+\w+.*?)\s+=+$'
    )
    # Regex untuk error/traceback
    error_section = False
    current_error = []
    current_test  = None

    i = 0
    while i < len(lines):
        line = lines[i]

        # Deteksi baris test result
        m = test_pattern.match(line)
        if m:
            file_path = m.group(1).replace('\\', '/')
            test_id   = m.group(2)
            status    = m.group(3).upper()

            # Parse module & test name dari test_id
            parts = test_id.split('::')
            if len(parts) >= 2:
                module_class = parts[0]
                test_name    = parts[-1]
            else:
                module_class = '-'
                test_name    = test_id

            # Ambil file name sebagai modul
            modul = file_path.split('/')[-1].replace('test_', '').replace('.py', '').replace('_', ' ').title()

            # Bersihkan nama test jadi deskripsi
            desc = test_name.replace('test_', '').replace('_', ' ')
            # Ambil ID (UT-AUTH-XX atau FT-XX)
            id_match = re.search(r'(UT_AUTH_\d+[a-z]?|FT_\w+_\d+[a-z]?)', test_name, re.IGNORECASE)
            test_case_id = id_match.group(1).replace('_', '-').upper() if id_match else '-'

            tests.append({
                'file'    : file_path,
                'modul'   : modul,
                'id'      : test_case_id,
                'nama'    : test_name,
                'deskripsi': desc,
                'status'  : status,
                'error'   : '',
            })
            current_test = tests[-1]
            error_section = False
            current_error = []

        # Deteksi error/traceback
        elif 'FAILED' in line and '::' in line and current_test:
            error_section = True

        elif error_section and current_test:
            if line.startswith('_') or line.startswith('='):
                if current_error:
                    current_test['error'] = '\n'.join(current_error[-8:])
                error_section = False
                current_error = []
            else:
                current_error.append(line)

        i += 1

    # Parse summary baris terakhir
    summary = {'passed': 0, 'failed': 0, 'error': 0, 'skipped': 0, 'total': 0}
    for line in reversed(lines):
        # Contoh: "15 passed, 3 failed, 1 error in 4.23s"
        nums = re.findall(r'(\d+)\s+(passed|failed|error|skipped)', line, re.IGNORECASE)
        if nums:
            for count, label in nums:
                key = label.lower()
                if key in summary:
                    summary[key] = int(count)
            break

    summary['total'] = sum(summary[v] for v in ['passed', 'failed', 'error', 'skipped'])

    # Kelompokkan per modul
    by_module = {}
    for t in tests:
        mod = t['modul']
        if mod not in by_module:
            by_module[mod] = []
        by_module[mod].append(t)

    return {
        'tests'     : tests,
        'by_module' : by_module,
        'summary'   : summary,
        'raw_output': output,
        'timestamp' : datetime.now().strftime('%d %B %Y, %H:%M WIB'),
        'date_short': datetime.now().strftime('%Y%m%d_%H%M'),
    }


# ══════════════════════════════════════════════════════════════
# STEP 3 — HELPER DOCX
# ══════════════════════════════════════════════════════════════

# Warna
CLR = {
    'primary'   : RGBColor(0x4F, 0x46, 0xE5),
    'success'   : RGBColor(0x06, 0x5F, 0x46),
    'success_bg': RGBColor(0xD1, 0xFA, 0xE5),
    'danger'    : RGBColor(0x99, 0x1B, 0x1B),
    'danger_bg' : RGBColor(0xFE, 0xE2, 0xE2),
    'warning'   : RGBColor(0x92, 0x40, 0x0E),
    'warning_bg': RGBColor(0xFE, 0xF3, 0xC7),
    'gray'      : RGBColor(0x47, 0x55, 0x69),
    'gray_lt'   : RGBColor(0xF1, 0xF5, 0xF9),
    'white'     : RGBColor(0xFF, 0xFF, 0xFF),
    'hdr_tbl'   : RGBColor(0x31, 0x2E, 0x81),
}

def hex_to_rgb_str(r, g, b):
    return f'{r:02X}{g:02X}{b:02X}'

def set_cell_bg(cell, hex_color: str):
    """Set background color pada cell tabel."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    """Set border tipis pada cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CBD5E1')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_cell_text(cell, text, bold=False, color=None, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Tambahkan teks ke cell dengan formatting."""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def make_header_row(table, headers, widths_cm, bg='312E81'):
    """Buat baris header tabel dengan styling."""
    row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, widths_cm)):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        add_cell_text(cell, hdr, bold=True, color=CLR['white'], size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

def add_data_row(table, values, widths_cm, bg='FFFFFF', colors=None, bolds=None):
    """Tambahkan baris data ke tabel."""
    row = table.add_row()
    colors = colors or [None] * len(values)
    bolds  = bolds  or [False] * len(values)
    for i, (val, w) in enumerate(zip(values, widths_cm)):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        add_cell_text(cell, val, bold=bolds[i], color=colors[i], size=9)


# ══════════════════════════════════════════════════════════════
# STEP 4 — BUAT DOKUMEN WORD
# ══════════════════════════════════════════════════════════════

def build_report(data: dict, output_path: str):
    """Buat dokumen Word laporan pengujian dari data parsed."""
    doc = Document()

    # ── Setup halaman ──────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(29.7)   # A4 landscape
    section.page_height   = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # ── Default style ──────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    summary = data['summary']
    tests   = data['tests']
    passed  = summary['passed']
    failed  = summary['failed']
    total   = summary['total'] or len(tests)
    pct     = round((passed / total * 100) if total > 0 else 0)

    # ════════════════════════════════════════════════════════════
    # HALAMAN JUDUL
    # ════════════════════════════════════════════════════════════

    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('LAPORAN HASIL PENGUJIAN SISTEM')
    r.font.size  = Pt(22)
    r.font.bold  = True
    r.font.color.rgb = CLR['primary']

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Modul Autentikasi & Verifikasi Email')
    r.font.size  = Pt(14)
    r.font.color.rgb = CLR['gray']

    app_name = doc.add_paragraph()
    app_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = app_name.add_run('Aplikasi SpeakUp — Public Speaking Trainer')
    r.font.size  = Pt(11)
    r.font.italic = True
    r.font.color.rgb = CLR['gray']

    doc.add_paragraph()

    # Info box tanggal & hasil
    info_tbl = doc.add_table(rows=1, cols=3)
    info_tbl.style = 'Table Grid'
    w_info = [8.3, 8.3, 8.3]

    cells_info = info_tbl.rows[0].cells

    # Tanggal
    set_cell_bg(cells_info[0], 'EEF2FF')
    set_cell_border(cells_info[0])
    p0 = cells_info[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f'📅 Tanggal Pengujian\n{data["timestamp"]}')
    r0.font.size = Pt(10)
    r0.font.color.rgb = CLR['primary']

    # Hasil
    result_bg  = 'D1FAE5' if failed == 0 else 'FEE2E2'
    result_clr = CLR['success'] if failed == 0 else CLR['danger']
    result_txt = '✅ SEMUA LULUS' if failed == 0 else f'⚠️ ADA {failed} GAGAL'
    set_cell_bg(cells_info[1], result_bg)
    set_cell_border(cells_info[1])
    p1 = cells_info[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f'Hasil Keseluruhan\n{result_txt}')
    r1.font.size  = Pt(10)
    r1.font.bold  = True
    r1.font.color.rgb = result_clr

    # Persentase
    set_cell_bg(cells_info[2], 'F0F4FF')
    set_cell_border(cells_info[2])
    p2 = cells_info[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f'Pass Rate\n{pct}%  ({passed}/{total})')
    r2.font.size  = Pt(10)
    r2.font.bold  = True
    r2.font.color.rgb = CLR['primary']

    for _ in range(8):
        doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # BAB 1 — RINGKASAN EKSEKUTIF
    # ════════════════════════════════════════════════════════════

    h1 = doc.add_heading('1. Ringkasan Eksekutif', level=1)
    h1.runs[0].font.color.rgb = CLR['primary']

    # Statistik ringkas
    stats_tbl = doc.add_table(rows=2, cols=5)
    stats_tbl.style = 'Table Grid'
    w_stats = [5.0, 5.0, 5.0, 5.0, 4.9]

    hdrs_stats = ['Total Test Case', 'PASSED ✅', 'FAILED ❌', 'ERROR ⚠️', 'SKIPPED ⏭']
    vals_stats = [
        str(total),
        str(summary['passed']),
        str(summary['failed']),
        str(summary.get('error', 0)),
        str(summary.get('skipped', 0)),
    ]
    bgs_stats  = ['312E81', '065F46', '991B1B', '92400E', '475569']
    bgs_val    = ['EEF2FF', 'D1FAE5', 'FEE2E2', 'FEF3C7', 'F1F5F9']

    for i, (h, v, bg_h, bg_v) in enumerate(zip(hdrs_stats, vals_stats, bgs_stats, bgs_val)):
        hc = stats_tbl.rows[0].cells[i]
        vc = stats_tbl.rows[1].cells[i]
        hc.width = Cm(w_stats[i])
        vc.width = Cm(w_stats[i])
        set_cell_bg(hc, bg_h)
        set_cell_bg(vc, bg_v)
        set_cell_border(hc)
        set_cell_border(vc)
        add_cell_text(hc, h, bold=True, color=CLR['white'], size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(vc, v, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Kesimpulan narasi
    p_kes = doc.add_paragraph()
    r_kes = p_kes.add_run('Kesimpulan: ')
    r_kes.font.bold = True
    r_kes.font.color.rgb = CLR['primary']

    if failed == 0 and summary.get('error', 0) == 0:
        kesimpulan = (
            f'Seluruh {total} test case berhasil dieksekusi dan lulus (PASSED) '
            f'dengan pass rate {pct}%. Modul autentikasi dan verifikasi email '
            f'berfungsi sesuai spesifikasi yang telah ditetapkan.'
        )
    else:
        kesimpulan = (
            f'Dari {total} test case yang dieksekusi, {passed} lulus dan '
            f'{failed + summary.get("error", 0)} gagal (pass rate {pct}%). '
            f'Test case yang gagal perlu ditindaklanjuti sebelum deployment.'
        )

    r_narasi = p_kes.add_run(kesimpulan)
    r_narasi.font.size  = Pt(10)
    r_narasi.font.color.rgb = CLR['gray']

    doc.add_paragraph()

    # ── Ringkasan per modul ────────────────────────────────────
    h2_mod = doc.add_heading('1.1 Ringkasan Per Modul', level=2)
    h2_mod.runs[0].font.color.rgb = CLR['gray']

    mod_tbl = doc.add_table(rows=1, cols=5)
    mod_tbl.style = 'Table Grid'
    w_mod = [5.5, 3.5, 3.0, 3.0, 9.9]
    make_header_row(mod_tbl,
        ['Modul / File Test', 'Total', 'Passed', 'Failed', 'Keterangan'],
        w_mod)

    for modul, mod_tests in data['by_module'].items():
        m_passed = sum(1 for t in mod_tests if t['status'] == 'PASSED')
        m_failed = sum(1 for t in mod_tests if t['status'] in ('FAILED', 'ERROR'))
        m_total  = len(mod_tests)
        ket      = '✅ Semua lulus' if m_failed == 0 else f'❌ {m_failed} gagal'
        row_bg   = 'F0FDF4' if m_failed == 0 else 'FFF7ED'

        add_data_row(mod_tbl,
            [modul, str(m_total), str(m_passed), str(m_failed), ket],
            w_mod, bg=row_bg,
            colors=[CLR['primary'], None, CLR['success'] if m_passed > 0 else None,
                    CLR['danger'] if m_failed > 0 else None, None],
            bolds=[True, True, True, True, False])

    doc.add_paragraph()
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # BAB 2 — DETAIL HASIL TEST PER MODUL
    # ════════════════════════════════════════════════════════════

    h1_detail = doc.add_heading('2. Detail Hasil Pengujian', level=1)
    h1_detail.runs[0].font.color.rgb = CLR['primary']

    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run(
        'Tabel berikut menampilkan hasil eksekusi setiap test case secara detail '
        'beserta status dan informasi error jika ada.'
    )
    r_intro.font.size  = Pt(10)
    r_intro.font.color.rgb = CLR['gray']

    doc.add_paragraph()

    STATUS_BG  = {'PASSED': 'D1FAE5', 'FAILED': 'FEE2E2', 'ERROR': 'FEF3C7',
                  'SKIPPED': 'F1F5F9', 'XFAIL': 'FEF3C7', 'XPASS': 'EEF2FF'}
    STATUS_CLR = {'PASSED': CLR['success'], 'FAILED': CLR['danger'],
                  'ERROR': CLR['warning'], 'SKIPPED': CLR['gray'],
                  'XFAIL': CLR['warning'], 'XPASS': CLR['primary']}
    STATUS_ICO = {'PASSED': '✅ PASSED', 'FAILED': '❌ FAILED', 'ERROR': '⚠️ ERROR',
                  'SKIPPED': '⏭ SKIPPED', 'XFAIL': 'XFAIL', 'XPASS': 'XPASS'}

    for idx, (modul, mod_tests) in enumerate(data['by_module'].items()):
        m_passed = sum(1 for t in mod_tests if t['status'] == 'PASSED')
        m_failed = sum(1 for t in mod_tests if t['status'] in ('FAILED', 'ERROR'))

        h2_m = doc.add_heading(f'2.{idx+1} {modul}', level=2)
        h2_m.runs[0].font.color.rgb = CLR['gray']

        p_stat = doc.add_paragraph()
        r_s1 = p_stat.add_run(f'Total: {len(mod_tests)}  |  ')
        r_s1.font.size = Pt(9)
        r_s2 = p_stat.add_run(f'Passed: {m_passed}  ')
        r_s2.font.size = Pt(9)
        r_s2.font.color.rgb = CLR['success']
        r_s2.font.bold = True
        r_s3 = p_stat.add_run(f'Failed: {m_failed}')
        r_s3.font.size = Pt(9)
        r_s3.font.color.rgb = CLR['danger'] if m_failed > 0 else CLR['gray']
        r_s3.font.bold = True

        # Tabel detail
        det_tbl = doc.add_table(rows=1, cols=5)
        det_tbl.style = 'Table Grid'
        w_det = [2.5, 5.5, 7.5, 2.5, 6.9]
        make_header_row(det_tbl,
            ['Test Case ID', 'Nama Fungsi', 'Deskripsi', 'Status', 'Catatan / Error'],
            w_det)

        for t in mod_tests:
            status     = t['status']
            status_ico = STATUS_ICO.get(status, status)
            status_clr = STATUS_CLR.get(status, CLR['gray'])
            status_bg  = STATUS_BG.get(status, 'FFFFFF')
            row_bg     = 'FFFFFF' if status == 'PASSED' else \
                         'FFF5F5' if status in ('FAILED', 'ERROR') else 'F8FAFC'

            error_note = t['error'][:120] + '...' if len(t['error']) > 120 else t['error']
            if not error_note and status == 'PASSED':
                error_note = '-'

            row = det_tbl.add_row()
            cells_d = row.cells

            # ID
            cells_d[0].width = Cm(w_det[0])
            set_cell_bg(cells_d[0], 'EEF2FF')
            set_cell_border(cells_d[0])
            add_cell_text(cells_d[0], t['id'], bold=True,
                          color=CLR['primary'], size=8)

            # Nama fungsi
            cells_d[1].width = Cm(w_det[1])
            set_cell_bg(cells_d[1], row_bg)
            set_cell_border(cells_d[1])
            add_cell_text(cells_d[1], t['nama'], size=8,
                          color=CLR['gray'])

            # Deskripsi
            cells_d[2].width = Cm(w_det[2])
            set_cell_bg(cells_d[2], row_bg)
            set_cell_border(cells_d[2])
            add_cell_text(cells_d[2], t['deskripsi'].replace('_', ' ').title(),
                          size=9)

            # Status
            cells_d[3].width = Cm(w_det[3])
            set_cell_bg(cells_d[3], status_bg)
            set_cell_border(cells_d[3])
            add_cell_text(cells_d[3], status_ico, bold=True,
                          color=status_clr, size=8,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

            # Catatan
            cells_d[4].width = Cm(w_det[4])
            set_cell_bg(cells_d[4], 'FFF5F5' if error_note != '-' else row_bg)
            set_cell_border(cells_d[4])
            add_cell_text(cells_d[4], error_note, size=8,
                          color=CLR['danger'] if error_note != '-' else CLR['gray'])

        doc.add_paragraph()

        # Jika ada yang gagal, tampilkan error lengkap
        failed_tests = [t for t in mod_tests if t['status'] in ('FAILED', 'ERROR') and t['error']]
        if failed_tests:
            h3_err = doc.add_heading(f'Detail Error — {modul}', level=3)
            h3_err.runs[0].font.color.rgb = CLR['danger']

            for t in failed_tests:
                p_err_title = doc.add_paragraph()
                r_et = p_err_title.add_run(f'❌  {t["id"]} — {t["nama"]}')
                r_et.font.bold  = True
                r_et.font.size  = Pt(9)
                r_et.font.color.rgb = CLR['danger']

                p_err_body = doc.add_paragraph()
                r_eb = p_err_body.add_run(t['error'])
                r_eb.font.size = Pt(8)
                r_eb.font.name = 'Courier New'
                r_eb.font.color.rgb = CLR['gray']
                p_err_body.paragraph_format.left_indent = Cm(0.5)

        if idx < len(data['by_module']) - 1:
            doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # BAB 3 — RAW OUTPUT PYTEST
    # ════════════════════════════════════════════════════════════

    doc.add_page_break()
    h1_raw = doc.add_heading('3. Raw Output Pytest', level=1)
    h1_raw.runs[0].font.color.rgb = CLR['primary']

    p_raw_desc = doc.add_paragraph()
    r_raw_desc = p_raw_desc.add_run(
        'Berikut adalah output lengkap dari eksekusi pytest (-v --tb=short):'
    )
    r_raw_desc.font.size = Pt(10)
    r_raw_desc.font.color.rgb = CLR['gray']

    doc.add_paragraph()

    # Tampilkan raw output dalam blok monospace
    raw_lines = data['raw_output'].splitlines()
    # Batasi 300 baris agar dokumen tidak terlalu besar
    if len(raw_lines) > 300:
        raw_lines = raw_lines[:300] + ['', '... (output dipotong, lihat file hasil_test.txt) ...']

    raw_para = doc.add_paragraph()
    raw_para.paragraph_format.left_indent = Cm(0.3)
    for line in raw_lines:
        if not line.strip():
            continue
        r_line = raw_para.add_run(line + '\n')
        r_line.font.name = 'Courier New'
        r_line.font.size = Pt(7.5)

        # Warnai baris PASSED / FAILED
        if 'PASSED' in line:
            r_line.font.color.rgb = CLR['success']
        elif 'FAILED' in line or 'ERROR' in line:
            r_line.font.color.rgb = CLR['danger']
        elif 'SKIPPED' in line:
            r_line.font.color.rgb = CLR['gray']
        elif line.startswith('=') or line.startswith('_'):
            r_line.font.bold = True
            r_line.font.color.rgb = CLR['primary']
        else:
            r_line.font.color.rgb = CLR['gray']

    # ════════════════════════════════════════════════════════════
    # BAB 4 — REKOMENDASI & TINDAK LANJUT
    # ════════════════════════════════════════════════════════════

    doc.add_page_break()
    h1_rek = doc.add_heading('4. Rekomendasi & Tindak Lanjut', level=1)
    h1_rek.runs[0].font.color.rgb = CLR['primary']

    if failed == 0 and summary.get('error', 0) == 0:
        items = [
            ('✅', 'Semua test case lulus.', 'Modul autentikasi siap untuk pengujian fungsional (black-box) dan deployment.'),
            ('🔄', 'Jalankan ulang secara berkala.', 'Integrasikan pytest ke pipeline CI/CD agar test otomatis berjalan setiap ada perubahan kode.'),
            ('📈', 'Tambah test coverage.', 'Pertimbangkan menambah test untuk edge case seperti SQL injection, brute force, dan rate limiting.'),
            ('🔒', 'Pengujian keamanan lanjutan.', 'Lakukan penetration testing pada endpoint login dan OAuth sebelum deployment ke production.'),
        ]
    else:
        items = [
            ('❌', f'Terdapat {failed + summary.get("error", 0)} test case gagal.', 'Perbaiki semua test yang FAILED sebelum melanjutkan ke tahap deployment.'),
            ('🔍', 'Analisis error.', 'Lihat detail error pada Bab 2 dan raw output pada Bab 3 untuk menentukan akar masalah.'),
            ('🔄', 'Re-test setelah perbaikan.', 'Jalankan ulang pytest setelah perbaikan untuk memastikan semua test lulus.'),
            ('📝', 'Dokumentasikan bug.', 'Catat setiap bug yang ditemukan di issue tracker sebelum diperbaiki.'),
        ]

    for icon, judul, detail in items:
        p_rek = doc.add_paragraph()
        r_icon = p_rek.add_run(f'{icon}  ')
        r_icon.font.size = Pt(11)
        r_judul = p_rek.add_run(f'{judul}  ')
        r_judul.font.bold  = True
        r_judul.font.size  = Pt(10)
        r_judul.font.color.rgb = CLR['primary']
        r_detail = p_rek.add_run(detail)
        r_detail.font.size = Pt(10)
        r_detail.font.color.rgb = CLR['gray']
        p_rek.paragraph_format.space_after = Pt(6)

    # ── Footer ─────────────────────────────────────────────────
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run(
        f'Dokumen ini di-generate otomatis oleh generate_test_report.py  •  '
        f'SpeakUp Testing Suite  •  {data["timestamp"]}'
    )
    r_footer.font.size  = Pt(8)
    r_footer.font.italic = True
    r_footer.font.color.rgb = CLR['gray']

    # ── Simpan ─────────────────────────────────────────────────
    doc.save(output_path)
    print(f'📄 Dokumen tersimpan: {output_path}')


# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    root = Path(__file__).parent

    # Jalankan pytest
    raw_output, returncode = run_pytest()

    # Simpan raw output ke TXT juga
    txt_path = root / 'tests' / 'hasil_test.txt'
    txt_path.write_text(raw_output, encoding='utf-8')
    print(f'📝 Raw output tersimpan: {txt_path}')

    # Parse output
    data = parse_pytest_output(raw_output)

    print(f'\n📊 Hasil:')
    print(f'   Total   : {data["summary"]["total"]}')
    print(f'   Passed  : {data["summary"]["passed"]}')
    print(f'   Failed  : {data["summary"]["failed"]}')
    print(f'   Error   : {data["summary"].get("error", 0)}')
    print(f'   Skipped : {data["summary"].get("skipped", 0)}')

    # Generate dokumen Word
    timestamp = data['date_short']
    docx_path = str(root / 'tests' / f'laporan_pengujian_{timestamp}.docx')
    build_report(data, docx_path)

    print('\n✅ Selesai!')
    print(f'   📄 Word  : tests/laporan_pengujian_{timestamp}.docx')
    print(f'   📝 TXT   : tests/hasil_test.txt')